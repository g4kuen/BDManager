from io import BytesIO

import openpyxl
from django.http import HttpResponse, HttpResponseNotFound, Http404, HttpResponseNotAllowed, HttpResponseRedirect
from django.shortcuts import render, redirect, get_object_or_404
from django.views.generic import UpdateView

from docx import Document


from .models import *
from .forms import *

menu = ['армия', 'cолдат', 'оружие', 'техника', 'коды баз', 'модель оружия' ,'тип техники']
title = 'главная страница'
def index(request):

    return render(request,'armia/index.html', {'title':title, 'menu': menu})
# Create your views here.


def categories(request, catid):
    if(request.GET):
        print(request.GET)
    return HttpResponse(f"<h1> категории </h1> <p>{catid}</p>")


def tables(request):
    table_name = request.GET.get('table_name')

    if table_name == 'EnginiryType':
        model_class = EnginiryType
    elif table_name == 'MilitaryBase':
        model_class = MilitaryBase
    elif table_name == 'WeaponModel':
        model_class = WeaponModel
    else:
        model_class = None

    records = model_class.objects.all() if model_class else []

    return render(request, 'armia/tables.html',
                  {'title': f'Table: {table_name}', 'table_name': table_name, 'records': records})


from django.http import HttpResponseRedirect
from django.urls import reverse

def add_item(request, table_name):
    if table_name == 'WeaponModel':
        form_class = WeaponModelForm
    elif table_name == 'EnginiryType':
        form_class = EnginiryTypeForm
    elif table_name == 'MilitaryBase':
        form_class = MilitaryBaseFormAdd
    elif table_name == 'Weapon':
        form_class = WeaponForm
    elif table_name == 'Soldier':
        form_class = SoldierForm
    elif table_name == 'BaseCodes':
        form_class = BaseCodesFormAdd
    elif table_name == 'Army':
        form_class = ArmyForm
    elif table_name == 'Enginiry':
        form_class = EnginiryForm
    else:
        form_class = None

    if request.method == 'POST':
        form = form_class(request.POST)
        if form.is_valid():
            form.save()
            return redirect('tables', table_name=table_name)
    else:
        form = form_class()

    return render(request, 'armia/add_item.html', {'form': form, 'table_name': table_name})


def edit_item(request, table_name, item_id):
    previous_id= 0
    if table_name == 'WeaponModel':
        model_class = WeaponModel
        form_class = WeaponModelForm
    elif table_name == 'EnginiryType':
        model_class = EnginiryType
        form_class = EnginiryTypeForm
    elif table_name == 'MilitaryBase':
        model_class = MilitaryBase
        form_class = MilitaryBaseForm
    elif table_name == 'Weapon':
        model_class = Weapon
        form_class = WeaponForm
    elif table_name == 'Soldier':
        model_class = Soldier
        form_class = SoldierForm
    elif table_name == 'BaseCodes':
        model_class = BaseCodes
        form_class = BaseCodesForm
    elif table_name == 'Army':
        model_class = Army
        form_class = ArmyForm
    elif table_name == 'Enginiry':
        model_class = Enginiry
        form_class = EnginiryForm
    else:
        return render(request, 'armia/error.html', {'error_message': 'Unknown entity'})

    item = get_object_or_404(model_class, pk=item_id)

    if request.method == 'POST':
        form = form_class(request.POST, instance=item)
        if form.is_valid():
            previous_id = item.id
            form.save()
            return redirect('tables', table_name=table_name)
    else:
        form = form_class(instance=item)

    return render(request, 'armia/redact.html', {'form': form, 'table_name': table_name, 'item_id': item_id, 'item': item, 'previous_id': previous_id})

#<a href="{% if item.id %}{% url 'edit_item' table_name=table_name item_id=item.pk %}{% endif %}">Редактировать</a>


from django.shortcuts import render, get_object_or_404, redirect
from .models import EnginiryType, WeaponModel

def delete_item(request, table_name, item_id):
    if table_name == 'WeaponModel':
        model_class = WeaponModel
    elif table_name == 'EnginiryType':
        model_class = EnginiryType
    elif table_name == 'MilitaryBase':
        model_class = MilitaryBase
    elif table_name == 'Weapon':
        model_class = Weapon
    elif table_name == 'Soldier':
        model_class = Soldier
    elif table_name == 'BaseCodes':
        model_class = BaseCodes
    elif table_name == 'Army':
        model_class = Army
    elif table_name == 'Enginiry':
        model_class = Enginiry
    else:
        return render(request, 'armia/error.html', {'error_message': 'Unknown entity'})

    item = get_object_or_404(model_class, pk=item_id)

    if request.method == 'POST':
        item.delete()
        return redirect('tables', table_name=table_name)

    return render(request, 'armia/delete_confirmation.html', {'item': item, 'table_url_name': table_name})

from django.shortcuts import render
from .models import WeaponModel, MilitaryBase, Weapon

def tables(request, table_name):
    if table_name == 'WeaponModel':
        model_class = WeaponModel
    elif table_name == 'MilitaryBase':
        model_class = MilitaryBase
    elif table_name == 'Weapon':
        model_class = Weapon
    elif table_name == 'EnginiryType':
        model_class = EnginiryType
    elif table_name == 'Soldier':
        model_class = Soldier
    elif table_name == 'BaseCodes':
        model_class = BaseCodes
    elif table_name == 'Army':
        model_class = Army
    elif table_name == 'Enginiry':
        model_class = Enginiry

    records = model_class.objects.all()

    return render(request, 'armia/tables.html', {'title': f'Table: {table_name}', 'menu': [], 'records': records, 'table_name': table_name})





def report_view(request):
    filter_value_report = request.GET.get('filter', '')
    filter_column_index_report = int(request.GET.get('filter_column', 0))
    print(filter_value_report)
    print(filter_column_index_report)
    filter_columns = ['full_name', 'rank', 'job', 'enginiry__function', 'weapon__type', 'weapon__ammo_amount',
                      'weapon__model__name', 'military_base__id', 'military_base__base_type']
    filter_column = filter_columns[filter_column_index_report]

    soldiers = Soldier.objects.filter(**{f'{filter_column}__icontains': filter_value_report})

    context = {
        'soldiers': soldiers,
        'filter_value_report':filter_value_report,
        'filter_column_index_report':filter_column_index_report
    }

    return render(request, 'armia/report.html', context)



def home_view(request):
    filter_value = request.GET.get('filter', '')
    soldiers = Soldier.objects.filter(full_name__icontains=filter_value)

    context = {
        'soldiers': soldiers,
    }

    return render(request, 'armia/home.html', context)



from concurrent.futures import ThreadPoolExecutor
def generate_excel_report(request):
    filter_value = request.GET.get('filter', '')
    print(request.GET.get('filter_column', 0), 'sadk,[pasdk[psakd[pddsadsadsa')
    filter_column_index = int(request.GET.get('filter_column', 0))
    filter_columns = ['full_name', 'rank', 'job', 'enginiry__function', 'weapon__type', 'weapon__ammo_amount',
                      'weapon__model__name', 'military_base__id', 'military_base__base_type']
    filter_column = filter_columns[filter_column_index]

    soldiers = Soldier.objects.filter(**{f'{filter_column}__icontains': filter_value})


    def generate_report_in_thread():

        workbook = openpyxl.Workbook()
        worksheet = workbook.active


        headers = ['Имя', 'Звание', 'Должность', 'техника','Тип оружия','Боезапас оружия','Модель оружия','локация базы','тип базы']
        worksheet.append(headers)


        for soldier in soldiers:
            row = [soldier.full_name, soldier.rank, soldier.job, str(soldier.enginiry.function), soldier.weapon.type, soldier.weapon.ammo_amount,str(soldier.weapon.model),soldier.military_base.id,soldier.military_base.base_type]
            worksheet.append(row)
        buffer = BytesIO()
        workbook.save(buffer)
        buffer.seek(0)

        return buffer.read()


    with ThreadPoolExecutor() as executor:
        report_content = executor.submit(generate_report_in_thread).result()


    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = 'attachment; filename=report.xlsx'
    response.write(report_content)

    return response


from django.shortcuts import render


def generate_docx_report(request):
    filter_value = request.GET.get('filter', '')
    print(request.GET.get('filter_column', 0),'sadk,[pasdk[psakd[pddsadsadsa')
    filter_column_index = int(request.GET.get('filter_column', 0))

    filter_columns = ['full_name', 'rank', 'job', 'enginiry__function', 'weapon__type', 'weapon__ammo_amount',
                      'weapon__model__name', 'military_base__id', 'military_base__base_type']
    filter_column = filter_columns[filter_column_index]

    soldiers = Soldier.objects.filter(**{f'{filter_column}__icontains': filter_value})
    print(soldiers,'czxzcxczxczxczx')
    def generate_report_in_thread():
        doc = Document()


        doc.add_heading('Отчет по солдатам', level=1)


        table = doc.add_table(rows=1, cols=9)
        headings = table.rows[0].cells
        headings[0].text = 'Имя'
        headings[1].text = 'Звание'
        headings[2].text = 'Должность'
        headings[3].text = 'Военная техника'
        headings[4].text = 'Тип Оружия'
        headings[5].text = 'Боезапас оружия'
        headings[6].text = 'Модель оружия'
        headings[7].text = 'локация военной базы'
        headings[8].text = 'тип военной базы'


        for soldier in soldiers:
            row_cells = table.add_row().cells
            row_cells[0].text = soldier.full_name
            row_cells[1].text = soldier.rank
            row_cells[2].text = soldier.job
            row_cells[3].text = str(soldier.enginiry)
            row_cells[4].text = str(soldier.weapon.type)
            row_cells[5].text = str(soldier.weapon.ammo_amount)
            row_cells[6].text = str(soldier.weapon.model)
            row_cells[7].text = str(soldier.military_base.id)
            row_cells[8].text = soldier.military_base.base_type


        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return buffer.read()


    with ThreadPoolExecutor() as executor:
        report_content = executor.submit(generate_report_in_thread).result()


    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=report.docx'
    response.write(report_content)

    return response

def get_model_and_form(table_name):
    models_forms = {
        'WeaponModel': (WeaponModel, WeaponModelForm),
        'EnginiryType': (EnginiryType, EnginiryTypeForm),
        'MilitaryBase': (MilitaryBase, MilitaryBaseForm),
        'Weapon': (Weapon, WeaponForm),
    }
    return models_forms.get(table_name, (None, None))

class EditItemView(UpdateView):
    template_name = 'armia/redact.html'
    fields = '__all__'

    def get_object(self, queryset=None):
        table_name = self.kwargs.get('table_name')
        item_id = self.kwargs.get('item_id')
        model_class, form_class = get_model_and_form(table_name)
        return get_object_or_404(model_class, pk=item_id)

    def get_success_url(self):
        table_name = self.kwargs.get('table_name')
        return reverse('tables', kwargs={'table_name': table_name})


def pageNotFound(request , exception):
    return HttpResponseNotFound('<h1>Страница не найдена</h1>')




